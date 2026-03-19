import os
import io
from datetime import date
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_excel(title, headers, rows):
    """通用Excel导出，返回BytesIO对象"""
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel sheet名最长31字符

    # 标题行
    title_font = Font(name="微软雅黑", size=14, bold=True)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    ws.cell(row=1, column=1, value=title).font = title_font
    ws.cell(row=1, column=1).alignment = Alignment(horizontal="center")

    # 表头
    header_font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    # 数据行
    data_font = Font(name="微软雅黑", size=10)
    for row_idx, row_data in enumerate(rows, 3):
        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = data_font
            cell.border = thin_border

    # 自动列宽
    for col in ws.columns:
        max_len = 0
        for cell in col:
            try:
                cell_len = len(str(cell.value or ""))
                if cell_len > max_len:
                    max_len = cell_len
            except Exception:
                pass
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def export_to_word(title, sections):
    """
    通用Word导出
    sections: [{"heading": "段落标题", "content": "内容"}, ...]
    或 sections: [{"heading": "标题", "table": {"headers": [...], "rows": [...]}}]
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)

    # 标题
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"导出日期：{date.today().isoformat()}")

    for sec in sections:
        heading = sec.get("heading", "")
        if heading:
            doc.add_heading(heading, level=2)

        if "content" in sec:
            doc.add_paragraph(sec["content"])

        if "table" in sec:
            tbl = sec["table"]
            headers = tbl["headers"]
            rows = tbl["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")

            # 表头
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True

            # 数据
            for row_idx, row_data in enumerate(rows):
                for col_idx, val in enumerate(row_data):
                    table.rows[row_idx + 1].cells[col_idx].text = str(val) if val is not None else ""

        if "paragraphs" in sec:
            for p in sec["paragraphs"]:
                doc.add_paragraph(p)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_record_to_word(title, subtitle, metadata_rows, content_sections, signatures=None, footer_text=""):
    """生成适合打印的单条记录Word文档。"""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.size = Pt(10.5)

    table = doc.add_table(rows=len(metadata_rows), cols=4, style="Table Grid")
    table.autofit = False
    for row_index, row_values in enumerate(metadata_rows):
        padded_values = list(row_values[:4])
        while len(padded_values) < 4:
            padded_values.append("")
        row_cells = table.rows[row_index].cells
        for cell_index, value in enumerate(padded_values):
            row_cells[cell_index].text = str(value or "")
            for paragraph in row_cells[cell_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    if cell_index in (0, 2):
                        run.bold = True

    doc.add_paragraph("")

    for section_item in content_sections:
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_item.get("heading", ""))
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        content = section_item.get("content", "") or ""
        for line in str(content).splitlines() or [""]:
            paragraph = doc.add_paragraph(line if line else " ")
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            paragraph.paragraph_format.line_spacing = 1.5

    signature_items = signatures or ["记录人签名", "审核人签名", "日期"]
    doc.add_paragraph("")
    signature_table = doc.add_table(rows=1, cols=len(signature_items))
    signature_table.autofit = True
    for index, label in enumerate(signature_items):
        cell = signature_table.rows[0].cells[index]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(f"{label}：{'_' * 14}")
        run.font.size = Pt(10.5)
        if index == len(signature_items) - 1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if footer_text:
        footer_paragraph = doc.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_paragraph.add_run(footer_text)
        footer_run.font.size = Pt(9)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
